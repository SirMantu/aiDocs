import tkinter as tk
import json
import os
from datetime import datetime
from tkinter import filedialog
from tkinter import messagebox
import openai  # Nutzung des OpenAI SDK
import docx  # Verwende python-docx für das Erstellen von Word-Dokumenten

# Default configuration file path
config_file = "config.json"

# Load configuration data from config.json if it exists
def load_config():
    if os.path.exists(config_file):
        with open(config_file, 'r') as file:
            return json.load(file)
    else:
        # Set default values if config doesn't exist
        return {
            'api_token': '',
            'default_directory': os.path.expanduser("~/Documents")
        }

# Save configuration data to config.json
def save_config(data):
    with open(config_file, 'w') as file:
        json.dump(data, file)

# Save the ChatGPT API token
def save_token():
    token = token_entry.get()
    if not token:
        messagebox.showerror("Fehler", "API-Token darf nicht leer sein.")
        return
    config_data = load_config()
    config_data['api_token'] = token
    save_config(config_data)
    messagebox.showinfo("Erfolg", "API-Token wurde gespeichert.")

# Load API token from configuration file
def load_token():
    config_data = load_config()
    return config_data.get('api_token', "")

# Select working directory for file saving/loading
def select_working_directory():
    global default_directory
    directory = filedialog.askdirectory(initialdir=default_directory, title="Arbeitsordner auswählen")
    if directory:
        default_directory = directory
        config_data = load_config()
        config_data['default_directory'] = default_directory
        save_config(config_data)
        messagebox.showinfo("Erfolg", f"Arbeitsordner wurde auf {default_directory} gesetzt.")

# Save patient data to JSON file
def save_data():
    vorname = vorname_entry.get()
    nachname = nachname_entry.get()
    
    if not vorname or not nachname:
        print("Vorname und Nachname müssen angegeben werden.")
        return

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{vorname}_{nachname}_{timestamp}.json"
    filepath = os.path.join(default_directory, filename)
    
    data = {
        "vorname": vorname,
        "nachname": nachname,
        "geschlecht": geschlecht_var.get(),
        "arzt": arzt_entry.get(),
        "notizen": chatgpt_output.get("1.0", tk.END),
        "diagnose": diagnose_box.get("1.0", tk.END),
        "arztbrief": arztbrief_box.get("1.0", tk.END),
        "chatgpt_eingabe": chatgpt_input.get("1.0", tk.END)
    }
    
    with open(filepath, 'w') as file:
        json.dump(data, file)
    print(f"Daten gespeichert in: {filepath}")

# Load patient data from JSON file
def load_data():
    file_path = filedialog.askopenfilename(
        title="Wähle eine Datei aus",
        initialdir=default_directory,
        filetypes=[("JSON Dateien", "*.json")]
    )
    
    if not file_path:
        return

    with open(file_path, 'r') as file:
        data = json.load(file)
        vorname_entry.delete(0, tk.END)
        vorname_entry.insert(0, data.get("vorname", ""))
        nachname_entry.delete(0, tk.END)
        nachname_entry.insert(0, data.get("nachname", ""))
        geschlecht_var.set(data.get("geschlecht", "Frau"))
        arzt_entry.delete(0, tk.END)
        arzt_entry.insert(0, data.get("arzt", ""))
        chatgpt_output.delete("1.0", tk.END)
        chatgpt_output.insert("1.0", data.get("notizen", ""))
        diagnose_box.delete("1.0", tk.END)
        diagnose_box.insert("1.0", data.get("diagnose", "Diagnose Notizen hier einfügen"))
        diagnose_box.config(fg="grey")
        chatgpt_input.delete("1.0", tk.END)
        chatgpt_input.insert("1.0", data.get("chatgpt_eingabe", "Ausgabe von ChatGPT"))
        chatgpt_input.config(fg="grey")

# Collect diagnosis input and send to ChatGPT
def collect_input_for_diagnosis():
    arztbrief = arztbrief_box.get("1.0", tk.END).strip()
    diagnose = diagnose_box.get("1.0", tk.END).strip()
    vorname = vorname_entry.get()
    nachname = nachname_entry.get()
    geschlecht = geschlecht_var.get()
    arzt = arzt_entry.get()
    
    # Erzeuge den Eingabetext
    input_text = (
        f"{arztbrief}\n{diagnose}\n"
        f"Informationen über den Patient:\n"
        f"Geschlecht: {geschlecht}\nVorname: {vorname}\nNachname: {nachname}\n"
        f"Behandelnder Arzt: {arzt}"
    )
    
    chatgpt_input.delete("1.0", tk.END)
    chatgpt_input.insert("1.0", input_text)

    # ChatGPT API Token laden
    api_token = load_token()
    if not api_token:
        messagebox.showerror("Fehler", "API-Token wurde nicht gefunden. Bitte füge den API-Token hinzu.")
        return

    # OpenAI API initialisieren
    openai.api_key = api_token

    try:
        # Anfrage an die GPT API senden (neue Schnittstelle)
        # response = openai.completions.create(
        #     model="gpt-3.5-turbo-0125",  # Verwende das gewünschte Modell
        #     prompt=input_text,
        #     max_tokens=1000,
        #     temperature=0.7
        # )
        
        # ChatGPT-Antwort anzeigen
        # chatgpt_reply = response['choices'][0]['text'].strip()
        # dummy text 
        # TODO DELETE after real use of chatgpd token
        chatgpt_reply = "Arztbrief mit Diagnose. BLABLABLA"
        chatgpt_output.delete("1.0", tk.END)
        chatgpt_output.insert("1.0", chatgpt_reply)

        # Antwort in ein Word-Dokument speichern
        timestamp = datetime.now().strftime("%Y_%m_%d")
        filename = f"{vorname}_{nachname}_{timestamp}.docx"
        filepath = os.path.join(default_directory, filename)

        doc = docx.Document()
        doc.add_heading(f"ChatGPT Diagnose für {vorname} {nachname}", 0)
        doc.add_paragraph(chatgpt_reply)
        doc.save(filepath)
        print(f"Antwort von ChatGPT gespeichert in: {filepath}")

    except Exception as e:
        messagebox.showerror("Fehler", f"Fehler bei der API-Anfrage: {e}")

# Handle focus events for placeholder text in Text widgets
def on_focus_in(event, widget, default_text):
    if widget.get("1.0", tk.END).strip() == default_text:
        widget.delete("1.0", tk.END)
        widget.config(fg="black")

def on_focus_out(event, widget, default_text):
    if not widget.get("1.0", tk.END).strip():
        widget.insert("1.0", default_text)
        widget.config(fg="grey")

# Hauptfenster erstellen
root = tk.Tk()
root.title("Patienteninformation")

# Load default configuration
config_data = load_config()
default_directory = config_data.get('default_directory', os.path.expanduser("~/Documents"))

# Textbox für Arztbrief
arztbrief_box = tk.Text(root, height=5, width=80)
arztbrief_box.grid(row=0, column=0, columnspan=2)
arztbrief_box.insert("1.0", "Schreibe einen Arztbrief mit den folgenden Diagnose Notizen")

# Textbox für Diagnose mit grauem Platzhaltertext
diagnose_box = tk.Text(root, height=10, width=80, fg="grey")
diagnose_box.grid(row=1, column=0, columnspan=2)
diagnose_box.insert("1.0", "Diagnose Notizen hier einfügen")
diagnose_box.bind("<FocusIn>", lambda event: on_focus_in(event, diagnose_box, "Diagnose Notizen hier einfügen"))
diagnose_box.bind("<FocusOut>", lambda event: on_focus_out(event, diagnose_box, "Diagnose Notizen hier einfügen"))

# Textbox für ChatGPT-Eingabe mit grauem Platzhaltertext
chatgpt_input = tk.Text(root, height=5, width=80, fg="grey")
chatgpt_input.grid(row=2, column=0, columnspan=2)
chatgpt_input.insert("1.0", "Ausgabe von ChatGPT")
chatgpt_input.bind("<FocusIn>", lambda event: on_focus_in(event, chatgpt_input, "Ausgabe von ChatGPT"))
chatgpt_input.bind("<FocusOut>", lambda event: on_focus_out(event, chatgpt_input, "Ausgabe von ChatGPT"))

# Textbox für ChatGPT-Ausgabe
chatgpt_output = tk.Text(root, height=10, width=80)
chatgpt_output.grid(row=3, column=0, columnspan=2)
chatgpt_output.insert("1.0", "Ausgabe von ChatGPT")
chatgpt_output.bind("<FocusIn>", lambda event: on_focus_in(event, chatgpt_output, "Ausgabe von ChatGPT"))
chatgpt_output.bind("<FocusOut>", lambda event: on_focus_out(event, chatgpt_output, "Ausgabe von ChatGPT"))

# Vorname
tk.Label(root, text="Vorname").grid(row=4, column=0)
vorname_entry = tk.Entry(root)
vorname_entry.grid(row=4, column=1)

# Nachname
tk.Label(root, text="Nachname").grid(row=5, column=0)
nachname_entry = tk.Entry(root)
nachname_entry.grid(row=5, column=1)

# Geschlecht als Radiobuttons
geschlecht_var = tk.StringVar(value="Frau")
geschlecht_frame = tk.LabelFrame(root, text="Geschlecht")
geschlecht_frame.grid(row=6, column=0, columnspan=2, padx=10, pady=5)
geschlecht_var = tk.StringVar(value="Frau")
tk.Radiobutton(geschlecht_frame, text="Frau", variable=geschlecht_var, value="Frau").pack(side="left", padx=10)
tk.Radiobutton(geschlecht_frame, text="Mann", variable=geschlecht_var, value="Mann").pack(side="left", padx=10)
tk.Radiobutton(geschlecht_frame, text="Divers", variable=geschlecht_var, value="Divers").pack(side="left", padx=10)


# Arztname
tk.Label(root, text="Behandelnder Arzt").grid(row=7, column=0)
arzt_entry = tk.Entry(root)
arzt_entry.grid(row=7, column=1)

# Save Data Button
save_button = tk.Button(root, text="Daten speichern", command=save_data)
save_button.grid(row=8, column=0)

# Load Data Button
load_button = tk.Button(root, text="Daten laden", command=load_data)
load_button.grid(row=8, column=1)

# Button zum Senden an ChatGPT
send_button = tk.Button(root, text="An ChatGPT senden", command=collect_input_for_diagnosis)
send_button.grid(row=9, column=0)

# Arbeitsordner auswählen
select_dir_button = tk.Button(root, text="Arbeitsordner auswählen", command=select_working_directory)
select_dir_button.grid(row=9, column=1)

# ChatGPT Token speichern
token_label = tk.Label(root, text="API-Token")
token_label.grid(row=10, column=0)
token_entry = tk.Entry(root)
token_entry.grid(row=10, column=1)
save_token_button = tk.Button(root, text="API-Token speichern", command=save_token)
save_token_button.grid(row=10, column=2)

root.mainloop()
